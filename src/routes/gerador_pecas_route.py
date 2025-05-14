# gerador_api_flask/src/routes/gerador_pecas_route.py
from flask import Blueprint, request, jsonify, send_file
import os
from docx import Document
from io import BytesIO
import logging # Adicionado para logging

gerador_pecas_bp = Blueprint("gerador_pecas_bp", __name__, url_prefix="/api/v1/gerador-pecas")

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "..", "templates_pecas")

def ensure_templates_dir():
    if not os.path.exists(TEMPLATES_DIR):
        os.makedirs(TEMPLATES_DIR)
        try:
            doc_template = Document()
            doc_template.add_heading("PETIÇÃO INICIAL", level=1)
            doc_template.add_paragraph("Autor: {{nome_autor}}")
            doc_template.add_paragraph("Réu: {{nome_reu}}")
            doc_template.add_paragraph("Objeto da Ação: {{objeto_acao}}")
            doc_template.add_paragraph("Valor da Causa: R$ {{valor_causa}}")
            doc_template.add_paragraph("\n[Local], [Data]")
            doc_template.add_paragraph("\n___________________________________")
            doc_template.add_paragraph("{{nome_advogado}}")
            doc_template.add_paragraph("OAB/UF {{oab_advogado}}")
            doc_template.save(os.path.join(TEMPLATES_DIR, "peticao_inicial_simples_template.docx"))
            logger.info(f"Template de exemplo criado em {os.path.join(TEMPLATES_DIR, 'peticao_inicial_simples_template.docx')}")
        except Exception as e:
            logger.error(f"Erro ao criar template de exemplo: {str(e)}")

# Função para substituir texto em parágrafos e runs
def replace_text_in_element(element, dados_variaveis):
    # Iterar para cada placeholder e seu valor
    for key, value in dados_variaveis.items():
        placeholder = f"{{{key}}}" # Ex: {{nome_autor}}
        str_value = str(value) # Garantir que o valor seja uma string

        # Lógica para substituir em parágrafos (element.text)
        if hasattr(element, 'text') and placeholder in element.text:
            # logger.info(f"Placeholder '{placeholder}' encontrado no texto do elemento. Substituindo por '{str_value}'. Texto original: '{element.text}'")
            element.text = element.text.replace(placeholder, str_value)
            # logger.info(f"Texto do elemento após substituição: '{element.text}'")

        # Lógica para substituir em runs (mais granular, pode ajudar com formatação)
        # Esta parte é mais complexa se o placeholder estiver dividido entre runs.
        # A abordagem abaixo substitui se o placeholder inteiro estiver em um run.
        if hasattr(element, 'runs'):
            for run in element.runs:
                if placeholder in run.text:
                    # logger.info(f"Placeholder '{placeholder}' encontrado no texto do run. Substituindo por '{str_value}'. Texto original do run: '{run.text}'")
                    run.text = run.text.replace(placeholder, str_value)
                    # logger.info(f"Texto do run após substituição: '{run.text}'")

@gerador_pecas_bp.route("/gerar-peca", methods=["POST"])
def gerar_peca_endpoint():
    ensure_templates_dir()
    data = request.get_json()
    if not data:
        logger.warning("Payload JSON ausente na requisição.")
        return jsonify({"error": "Payload JSON ausente"}), 400

    template_id = data.get("template_id")
    dados_variaveis = data.get("dados_variaveis")
    formato_saida = data.get("formato_saida", "docx")

    logger.info(f"Requisição recebida: template_id='{template_id}', dados_variaveis='{dados_variaveis}', formato_saida='{formato_saida}'")

    if not template_id or not isinstance(dados_variaveis, dict):
        logger.warning("template_id obrigatório ou dados_variaveis não é um dicionário.")
        return jsonify({"error": "template_id é obrigatório e dados_variaveis deve ser um objeto JSON."}), 400

    if template_id == "peticao_inicial_simples":
        template_filename = "peticao_inicial_simples_template.docx"
        template_path = os.path.join(TEMPLATES_DIR, template_filename)
        
        if not os.path.exists(template_path):
            logger.error(f"Template '{template_filename}' não encontrado em '{template_path}'.")
            return jsonify({"error": f"Template '{template_filename}' não encontrado."}), 404

        try:
            doc = Document(template_path)
            
            logger.info("Iniciando substituição em parágrafos do corpo...")
            for p in doc.paragraphs:
                replace_text_in_element(p, dados_variaveis)
            
            logger.info("Iniciando substituição em tabelas...")
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p_in_cell in cell.paragraphs:
                            replace_text_in_element(p_in_cell, dados_variaveis)
            
            # Adicionar substituição em cabeçalhos e rodapés (se necessário e suportado)
            # A biblioteca python-docx tem limitações para editar cabeçalhos/rodapés complexos.
            # Para substituições simples em cabeçalhos/rodapés:
            logger.info("Iniciando substituição em cabeçalhos...")
            for section in doc.sections:
                if section.header is not None:
                    for p_in_header in section.header.paragraphs:
                        replace_text_in_element(p_in_header, dados_variaveis)
                if section.footer is not None:
                    for p_in_footer in section.footer.paragraphs:
                         replace_text_in_element(p_in_footer, dados_variaveis)

            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)
            logger.info(f"Documento '{template_id}_gerada.docx' processado e pronto para envio.")

            if formato_saida == "docx":
                return send_file(
                    file_stream,
                    as_attachment=True,
                    download_name=f"{template_id}_gerada.docx",
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                logger.warning(f"Formato de saída '{formato_saida}' não suportado.")
                return jsonify({"error": "Formato de saída não suportado no momento."}), 400

        except Exception as e:
            logger.error(f"Erro ao gerar o documento: {str(e)}", exc_info=True)
            return jsonify({"error": f"Erro interno ao gerar o documento: {str(e)}"}), 500
    else:
        logger.warning(f"template_id '{template_id}' inválido ou não suportado.")
        return jsonify({"error": "template_id inválido ou não suportado no momento"}), 400

