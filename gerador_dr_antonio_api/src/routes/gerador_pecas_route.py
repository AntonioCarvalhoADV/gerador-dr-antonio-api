# gerador_api_flask/src/routes/gerador_pecas_route.py
from flask import Blueprint, request, jsonify, send_file
import os
from docx import Document
from io import BytesIO

gerador_pecas_bp = Blueprint("gerador_pecas_bp", __name__, url_prefix="/api/v1/gerador-pecas")

# O diretório de templates será relativo à raiz do projeto Flask quando empacotado.
# Vamos assumir que os templates estarão em 'src/templates_pecas' dentro do projeto Flask.
TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "..", "templates_pecas")

# Função auxiliar para garantir que o diretório de templates exista
def ensure_templates_dir():
    if not os.path.exists(TEMPLATES_DIR):
        os.makedirs(TEMPLATES_DIR)
        # Criar um template de exemplo se o diretório for criado agora
        try:
            doc_template = Document()
            doc_template.add_heading("PETIÇÃO INICIAL", level=1)
            doc_template.add_paragraph("Autor: {{nome_autor}}")
            doc_template.add_paragraph("Réu: {{nome_reu}}")
            doc_template.add_paragraph("Objeto da Ação: {{objeto_acao}}")
            doc_template.add_paragraph("Valor da Causa: R$ {{valor_causa}}")
            doc_template.add_paragraph("\n[Local], [Data]")
            doc_template.add_paragraph("\n___________________________________")
            doc_template.add_paragraph("{{nome_advogado}}")
            doc_template.add_paragraph("OAB/UF {{oab_advogado}}")
            doc_template.save(os.path.join(TEMPLATES_DIR, "peticao_inicial_simples_template.docx"))
        except Exception as e:
            print(f"Erro ao criar template de exemplo: {str(e)}") # Logar erro

@gerador_pecas_bp.route("/gerar-peca", methods=["POST"])
def gerar_peca_endpoint():
    ensure_templates_dir() # Garante que o diretório e o template de exemplo existam
    data = request.get_json()
    if not data:
        return jsonify({"error": "Payload JSON ausente"}), 400

    template_id = data.get("template_id")
    dados_variaveis = data.get("dados_variaveis")
    formato_saida = data.get("formato_saida", "docx")

    if not template_id or not dados_variaveis:
        return jsonify({"error": "template_id e dados_variaveis são obrigatórios"}), 400

    if template_id == "peticao_inicial_simples":
        template_filename = "peticao_inicial_simples_template.docx"
        template_path = os.path.join(TEMPLATES_DIR, template_filename)
        
        if not os.path.exists(template_path):
             return jsonify({"error": f"Template \'{template_filename}\' não encontrado. O diretório de templates é: {TEMPLATES_DIR}"}), 404

        try:
            doc = Document(template_path)
            for p in doc.paragraphs:
                for key, value in dados_variaveis.items():
                    placeholder = f"{{{{{key}}}}}"
                    if placeholder in p.text:
                        # Esta é uma substituição básica. Para cenários mais complexos,
                        # pode ser necessário iterar sobre p.runs e substituir dentro de cada run.
                        p.text = p.text.replace(placeholder, str(value))
            
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            if formato_saida == "docx":
                return send_file(
                    file_stream,
                    as_attachment=True,
                    download_name=f"{template_id}_gerada.docx",
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                return jsonify({"error": "Formato de saída não suportado no momento."}), 400

        except Exception as e:
            return jsonify({"error": f"Erro ao gerar o documento: {str(e)}"}), 500
    else:
        return jsonify({"error": "template_id inválido ou não suportado no momento"}), 400

